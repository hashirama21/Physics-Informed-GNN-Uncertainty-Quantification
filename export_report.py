#!/usr/bin/env python3
"""
export_report.py — Convert results.md to DOCX and/or PDF.

Usage:
    python export_report.py                    # → results.docx + results.pdf
    python export_report.py --docx             # → results.docx only
    python export_report.py --pdf              # → results.pdf only
    python export_report.py --input my.md      # custom input file

Priority order:
    1. pandoc  (best quality — installs via 'brew install pandoc' on macOS)
    2. Pure Python fallback (auto-installs: markdown, python-docx, weasyprint)
"""

from __future__ import annotations

import argparse
import subprocess
import sys
import textwrap
from pathlib import Path

ROOT = Path(__file__).resolve().parent
DEFAULT_INPUT  = ROOT / "results.md"
DEFAULT_DOCX   = ROOT / "results.docx"
DEFAULT_PDF    = ROOT / "results.pdf"


def _pandoc_available() -> bool:
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False


def _install_pandoc_macos() -> bool:
    """Try to install pandoc via brew on macOS."""
    try:
        subprocess.run(["brew", "install", "pandoc"], check=True)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False


def pandoc_to_docx(src: Path, dst: Path) -> bool:
    """Convert markdown → DOCX via pandoc."""
    cmd = [
        "pandoc", str(src),
        "-o", str(dst),
        "--from", "markdown+pipe_tables+fenced_code_blocks",
        "--to", "docx",
        "--standalone",
        "--highlight-style", "tango",
        "-V", "geometry:margin=2.5cm",
    ]
    ref = ROOT / "reference.docx"
    if ref.exists():
        cmd += ["--reference-doc", str(ref)]
    try:
        subprocess.run(cmd, check=True)
        return True
    except subprocess.CalledProcessError as exc:
        print(f"[pandoc DOCX] error: {exc}", file=sys.stderr)
        return False


def pandoc_to_pdf(src: Path, dst: Path) -> bool:
    """Convert markdown → PDF via pandoc (requires LaTeX or wkhtmltopdf)."""
    # Try with wkhtmltopdf engine first (no LaTeX needed)
    for engine in ["wkhtmltopdf", "weasyprint", None]:
        cmd = [
            "pandoc", str(src),
            "-o", str(dst),
            "--from", "markdown+pipe_tables+fenced_code_blocks",
            "--to", "pdf",
            "--pdf-engine-opt=--enable-local-file-access",
            "-V", "geometry:margin=2.5cm",
            "-V", "fontsize=11pt",
            "-V", "colorlinks=true",
        ]
        if engine:
            cmd += ["--pdf-engine", engine]
        try:
            subprocess.run(cmd, check=True, capture_output=True)
            return True
        except subprocess.CalledProcessError:
            continue
    print("[pandoc PDF] no PDF engine available; trying Python fallback.", file=sys.stderr)
    return False

def _pip_install(*packages: str) -> None:
    subprocess.run(
        [sys.executable, "-m", "pip", "install", "--quiet", *packages],
        check=True,
    )


def _ensure_deps() -> None:
    missing = []
    for pkg, imp in [("markdown", "markdown"), ("python-docx", "docx"),
                     ("weasyprint", "weasyprint")]:
        try:
            __import__(imp)
        except ImportError:
            missing.append(pkg)
    if missing:
        print(f"[setup] installing: {', '.join(missing)} …")
        _pip_install(*missing)

_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:wght@400;600;700&family=Source+Code+Pro&display=swap');

body {
    font-family: 'Source Sans Pro', 'Helvetica Neue', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #1a1a2e;
    max-width: 850px;
    margin: 0 auto;
    padding: 2cm 2.5cm;
}
h1 { font-size: 22pt; color: #16213e; border-bottom: 2px solid #0f3460; padding-bottom: 6px; margin-top: 2em; }
h2 { font-size: 16pt; color: #0f3460; border-bottom: 1px solid #e0e0e0; padding-bottom: 4px; margin-top: 1.8em; }
h3 { font-size: 13pt; color: #533483; margin-top: 1.4em; }
h4 { font-size: 11pt; color: #444; margin-top: 1.2em; }
table {
    border-collapse: collapse;
    width: 100%;
    margin: 1em 0;
    font-size: 9.5pt;
    page-break-inside: avoid;
}
th {
    background: #0f3460;
    color: white;
    padding: 6px 10px;
    text-align: left;
    font-weight: 600;
}
td {
    padding: 5px 10px;
    border-bottom: 1px solid #e5e5e5;
}
tr:nth-child(even) td { background: #f8f9fb; }
tr:hover td { background: #eef2ff; }
code {
    font-family: 'Source Code Pro', 'Courier New', monospace;
    font-size: 9pt;
    background: #f4f4f8;
    padding: 1px 4px;
    border-radius: 3px;
    color: #c0392b;
}
pre {
    background: #1a1a2e;
    color: #e8e8e8;
    padding: 14px 18px;
    border-radius: 6px;
    font-size: 8.5pt;
    line-height: 1.5;
    overflow-x: auto;
    page-break-inside: avoid;
}
pre code {
    background: transparent;
    color: #e8e8e8;
    padding: 0;
}
blockquote {
    border-left: 4px solid #0f3460;
    margin: 1em 0;
    padding: 0.5em 1em;
    color: #555;
    background: #f8f9fb;
}
hr { border: none; border-top: 2px solid #e0e0e0; margin: 2em 0; }
strong { color: #0f3460; }
@page {
    margin: 2cm 2.5cm;
    @bottom-center {
        content: "PIGNN-UQ — DONGMO — " counter(page) " / " counter(pages);
        font-size: 9pt;
        color: #888;
    }
}
"""


def python_to_docx(src: Path, dst: Path) -> bool:
    """markdown → python-docx DOCX (simplified structural conversion)."""
    try:
        import markdown as md_lib
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        _ensure_deps()
        import markdown as md_lib
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches

    text = src.read_text(encoding="utf-8")
    doc  = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    in_code_block = False
    code_lines: list[str] = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        code_lines.clear()

    in_table      = False
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        # Filter out separator rows (---|---|...)
        data = [r for r in table_rows if not all(c.strip().replace("-","").replace(":","") == "" for c in r)]
        if not data:
            table_rows.clear()
            return
        n_cols = max(len(r) for r in data)
        tbl = doc.add_table(rows=len(data), cols=n_cols)
        tbl.style = "Table Grid"
        for ri, row in enumerate(data):
            for ci in range(n_cols):
                cell = tbl.cell(ri, ci)
                val  = row[ci].strip() if ci < len(row) else ""
                cell.text = val
                if ri == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
        doc.add_paragraph()
        table_rows.clear()

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Code block ──────────────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code_block:
                flush_table()
                in_code_block = True
            else:
                flush_code()
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("|"):
            cols = [c for c in line.split("|") if c != ""]
            table_rows.append(cols)
            in_table = True
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False

        stripped = line.strip()
        if stripped.startswith("#### "):
            flush_table()
            p = doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            flush_table()
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            flush_table()
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            flush_table()
            p = doc.add_heading(stripped[2:], level=1)

        elif stripped.startswith("---"):
            flush_table()
            doc.add_paragraph("─" * 60)

        elif stripped.startswith("> "):
            flush_table()
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True

        elif stripped.startswith("- ") or stripped.startswith("* "):
            flush_table()
            doc.add_paragraph(stripped[2:], style="List Bullet")

        elif stripped:
            flush_table()
            doc.add_paragraph(stripped)

        else:
            flush_table()

        i += 1

    flush_code()
    flush_table()
    doc.save(str(dst))
    return True

def python_to_pdf(src: Path, dst: Path) -> bool:
    """markdown → HTML → PDF via weasyprint."""
    try:
        import markdown as md_lib
        import weasyprint
    except ImportError:
        _ensure_deps()
        import markdown as md_lib
        import weasyprint  # type: ignore

    text = src.read_text(encoding="utf-8")
    body = md_lib.markdown(
        text,
        extensions=["tables", "fenced_code", "nl2br", "toc"],
    )
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="utf-8">
  <title>PIGNN-UQ — Training Log</title>
  <style>{_CSS}</style>
</head>
<body>
{body}
</body>
</html>"""

    html_path = dst.with_suffix(".html")
    html_path.write_text(html, encoding="utf-8")

    try:
        weasyprint.HTML(filename=str(html_path)).write_pdf(str(dst))
        html_path.unlink()
        return True
    except Exception as exc:
        print(f"[weasyprint] error: {exc}", file=sys.stderr)
        print(f"[weasyprint] HTML saved at: {html_path}", file=sys.stderr)
        print("[tip] Open the HTML in a browser and print to PDF as fallback.", file=sys.stderr)
        return False


def export_html(src: Path, dst: Path) -> bool:
    """markdown → styled HTML (no extra deps — pure stdlib + markdown)."""
    try:
        import markdown as md_lib
    except ImportError:
        _pip_install("markdown")
        import markdown as md_lib  # type: ignore

    text = src.read_text(encoding="utf-8")
    body = md_lib.markdown(
        text,
        extensions=["tables", "fenced_code", "nl2br", "toc"],
    )
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="utf-8">
  <title>PIGNN-UQ — Training Log</title>
  <style>{_CSS}</style>
</head>
<body>
{body}
</body>
</html>"""
    dst.write_text(html, encoding="utf-8")
    return True

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Export results.md to DOCX and/or PDF.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
        Examples:
          python export_report.py                    # DOCX + PDF
          python export_report.py --docx             # DOCX only
          python export_report.py --pdf              # PDF only
          python export_report.py --html             # HTML only (always works)
          python export_report.py --input my.md      # custom input
        """),
    )
    parser.add_argument("--input",  default=str(DEFAULT_INPUT),  help="Source .md file")
    parser.add_argument("--output", default=None,                help="Output base path (no extension)")
    parser.add_argument("--docx",   action="store_true",         help="Export DOCX")
    parser.add_argument("--pdf",    action="store_true",         help="Export PDF")
    parser.add_argument("--html",   action="store_true",         help="Export HTML")
    args = parser.parse_args()

    src = Path(args.input)
    if not src.exists():
        sys.exit(f"Input file not found: {src}")

    base   = Path(args.output) if args.output else src.with_suffix("")
    do_all = not (args.docx or args.pdf or args.html)

    # ── Detect pandoc ──────────────────────────────────────────────────────
    has_pandoc = _pandoc_available()
    if not has_pandoc and sys.platform == "darwin":
        print("[setup] pandoc not found — trying 'brew install pandoc' …")
        if _install_pandoc_macos():
            has_pandoc = _pandoc_available()
            if has_pandoc:
                print("[setup] pandoc installed successfully.")

    # ── DOCX ───────────────────────────────────────────────────────────────
    if args.docx or do_all:
        dst = base.with_suffix(".docx")
        print(f"[DOCX] {src.name} → {dst.name}")
        ok = False
        if has_pandoc:
            ok = pandoc_to_docx(src, dst)
        if not ok:
            print("[DOCX] falling back to python-docx …")
            _ensure_deps()
            ok = python_to_docx(src, dst)
        print(f"[DOCX] {'✓ done' if ok else '✗ failed'} → {dst}")

    # ── PDF ────────────────────────────────────────────────────────────────
    if args.pdf or do_all:
        dst = base.with_suffix(".pdf")
        print(f"[PDF]  {src.name} → {dst.name}")
        ok = False
        if has_pandoc:
            ok = pandoc_to_pdf(src, dst)
        if not ok:
            print("[PDF]  falling back to weasyprint …")
            _ensure_deps()
            ok = python_to_pdf(src, dst)
        print(f"[PDF]  {'✓ done' if ok else '✗ failed (see HTML fallback above)'} → {dst}")

    # ── HTML ───────────────────────────────────────────────────────────────
    if args.html or (do_all and False):  # HTML only when explicitly requested
        dst = base.with_suffix(".html")
        print(f"[HTML] {src.name} → {dst.name}")
        ok  = export_html(src, dst)
        print(f"[HTML] {'✓ done' if ok else '✗ failed'} → {dst}")

    # ── Tip: HTML fallback is always safe ─────────────────────────────────
    if do_all or args.pdf:
        html_dst = base.with_suffix(".html")
        if not html_dst.exists():
            print("\n[tip]  Run with --html to get a styled HTML you can print-to-PDF from any browser.")


if __name__ == "__main__":
    main()
